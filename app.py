from flask import Flask, render_template, request, jsonify, send_file
import os
from dotenv import load_dotenv
import requests
from PyPDF2 import PdfReader
from docx import Document
import io
import tempfile

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Available languages dictionary
LANGUAGES = {
    '1': {'code': 'en', 'name': 'English'},
    '2': {'code': 'es', 'name': 'Spanish'},
    '3': {'code': 'fr', 'name': 'French'},
    '4': {'code': 'de', 'name': 'German'},
    '5': {'code': 'pt', 'name': 'Portuguese'}
}

def translate_text(text, target_language):
    """Translate text to the target language."""
    try:
        api_key = os.getenv('GOOGLE_API_KEY')
        if not api_key:
            return "Error: GOOGLE_API_KEY not found in .env file"
            
        url = f"https://translation.googleapis.com/language/translate/v2?key={api_key}"
        data = {
            'q': text,
            'target': target_language
        }
        
        response = requests.post(url, json=data)
        if response.status_code == 200:
            result = response.json()
            return result['data']['translations'][0]['translatedText']
        else:
            return f"Translation failed with status code: {response.status_code}"
            
    except Exception as e:
        return f"Error during translation: {str(e)}"

def extract_text_from_pdf(file):
    """Extract text from PDF file."""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file):
    """Extract text from DOCX file."""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def create_translated_docx(text, target_language):
    """Create a new DOCX file with translated text."""
    try:
        translated_text = translate_text(text, target_language)
        doc = Document()
        doc.add_paragraph(translated_text)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
    except Exception as e:
        return f"Error creating translated DOCX: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html', languages=LANGUAGES)

@app.route('/translate', methods=['POST'])
def translate():
    data = request.get_json()
    text = data.get('text', '')
    target_language = data.get('target_language', 'en')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    translated_text = translate_text(text, target_language)
    return jsonify({'translated_text': translated_text})

@app.route('/translate_file', methods=['POST'])
def translate_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    target_language = request.form.get('target_language', 'en')
    
    if not file:
        return jsonify({'error': 'No file selected'}), 400
    
    # Get file extension
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    try:
        if file_extension == '.pdf':
            text = extract_text_from_pdf(file)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file)
        else:
            return jsonify({'error': 'Unsupported file format'}), 400
        
        if text.startswith('Error'):
            return jsonify({'error': text}), 400
        
        # Create translated document
        translated_file = create_translated_docx(text, target_language)
        if translated_file.startswith('Error'):
            return jsonify({'error': translated_file}), 400
        
        # Send the file
        return send_file(
            translated_file,
            as_attachment=True,
            download_name=f'translated_{file.filename}',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Get port from environment variable or use default
    port = int(os.environ.get('PORT', 5000))
    # Run the app
    app.run(host='0.0.0.0', port=port) 